#!/usr/bin/env python3
from __future__ import annotations

import argparse
from html.parser import HTMLParser
import json
import re
import sys
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


CURRENT_DIR = Path(__file__).resolve().parent
TOOLS_DIR = CURRENT_DIR.parent.parent
LAYOUT_REBUILD_DIR = TOOLS_DIR / "vl_layout_rebuild"

for import_dir in (CURRENT_DIR, LAYOUT_REBUILD_DIR):
    import_path = str(import_dir)
    if import_path not in sys.path:
        sys.path.insert(0, import_path)

from pipeline import (  # noqa: E402
    Block,
    _build_hybrid_blocks,
    _build_parsing_blocks,
    _choose_blocks_auto,
    _choose_generic_coord_size,
    _extract_generic_pages,
    _extract_mineru_pages,
    _extract_pages,
    _get_input_image_size,
    _resolve_image_bytes,
    _restore_form_layout,
)
from translator import Translator  # noqa: E402
from vl_layout_to_docx import _char_visual_unit, add_image_box, add_textbox  # noqa: E402


# 临时仅保留三路 OCR：paddleocr-vl / mineru / dotsocr
# ENGINES = ("paddleocr-vl", "mineru", "dotsocr", "opendataloader-hybrid-full")
ENGINES = ("paddleocr-vl", "mineru", "dotsocr")
VISUAL_ONLY_LABELS = {
    "picture",
    "image",
    "image_body",
    "figure",
    "chart",
    "table",
    "table_body",
    "formula",
    "seal",
}

IMAGE_LIKE_LABELS = {"image", "image_body", "picture", "figure", "chart"}


class _TableHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.rows: list[list[dict[str, Any]]] = []
        self._current_row: list[dict[str, Any]] | None = None
        self._current_cell: dict[str, Any] | None = None
        self._in_cell = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag == "tr":
            self._current_row = []
        elif tag in {"td", "th"} and self._current_row is not None:
            attr_map = {k.lower(): v for k, v in attrs if k}
            self._current_cell = {
                "text": "",
                "rowspan": int(attr_map.get("rowspan") or 1),
                "colspan": int(attr_map.get("colspan") or 1),
                "header": tag == "th",
            }
            self._in_cell = True

    def handle_data(self, data: str) -> None:
        if self._in_cell and self._current_cell is not None:
            self._current_cell["text"] += data

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"td", "th"} and self._current_row is not None and self._current_cell is not None:
            self._current_cell["text"] = " ".join(str(self._current_cell["text"]).split())
            self._current_row.append(self._current_cell)
            self._current_cell = None
            self._in_cell = False
        elif tag == "tr" and self._current_row is not None:
            self.rows.append(self._current_row)
            self._current_row = None


def _parse_html_table(html: str) -> list[list[dict[str, Any]]]:
    if "<table" not in (html or "").lower():
        return []
    parser = _TableHTMLParser()
    try:
        parser.feed(html)
    except Exception:
        return []
    return [row for row in parser.rows if row]


def _table_rows_from_opendataloader_item(item: dict[str, Any]) -> list[list[dict[str, Any]]]:
    rows = item.get("rows")
    if not isinstance(rows, list):
        return []
    table_rows: list[list[dict[str, Any]]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        cells = row.get("cells")
        if not isinstance(cells, list):
            continue
        out_row: list[dict[str, Any]] = []
        for cell in cells:
            if not isinstance(cell, dict):
                continue
            kids = cell.get("kids")
            texts: list[str] = []
            if isinstance(kids, list):
                for kid in kids:
                    if isinstance(kid, dict):
                        text = str(kid.get("content") or "").strip()
                        if text:
                            texts.append(text)
            out_row.append(
                {
                    "text": "\n".join(texts).strip(),
                    "rowspan": max(1, int(cell.get("row span") or 1)),
                    "colspan": max(1, int(cell.get("column span") or 1)),
                    "header": False,
                }
            )
        if out_row:
            table_rows.append(out_row)
    return table_rows


def _flip_bbox_vertical(bbox: list[float] | tuple[float, float, float, float], page_height: float) -> list[float]:
    x1, y1, x2, y2 = [float(v) for v in bbox]
    return [x1, max(0.0, page_height - y2), x2, max(0.0, page_height - y1)]


def _opendataloader_uses_bottom_left_origin(page_item: dict[str, Any]) -> bool:
    raw = _unwrap_raw_meta(page_item.get("raw"))
    if not isinstance(raw, dict):
        return False
    return bool(raw.get("_opendataloader_bottom_left_origin"))


def _clone_block_with_bbox(block: Block, bbox: list[float]) -> Block:
    return Block(
        bbox=bbox,
        label=block.label,
        order=block.order,
        text=block.text,
        image_ref=block.image_ref,
    )


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export DOCX test files from vl_visual_clone raw JSON outputs")
    parser.add_argument(
        "--input-dir",
        default="/u01/huzekun/data/FUNSD/test",
        help="Directory that contains *_<engine>_raw.json files",
    )
    parser.add_argument(
        "--output-dir",
        default="/u01/huzekun/data/FUNSD/test_docx",
        help="Directory to store generated DOCX files",
    )
    parser.add_argument(
        "--pdf-search-dir",
        action="append",
        default=[
            "/u01/huzekun/data/FUNSD/pdf_test",
            "/u01/huzekun/data/FUNSD/data",
        ],
        help="Directory used to resolve original input PDFs; can be passed multiple times",
    )
    parser.add_argument(
        "--geometry-source",
        choices=["auto", "parsing", "layout-hybrid"],
        default="auto",
        help="Geometry strategy aligned with visual clone pipeline",
    )
    parser.add_argument(
        "--with-page-background",
        action="store_true",
        help="Render original PDF page as a DOCX background image",
    )
    parser.add_argument(
        "--font-name",
        default="SimSun",
        help="DOCX font name",
    )
    parser.add_argument(
        "--font-size",
        type=float,
        default=10.0,
        help="Default font size in pt",
    )
    parser.add_argument(
        "--min-font-size",
        type=float,
        default=4.0,
        help="Minimal font size in pt",
    )
    parser.add_argument(
        "--text-box-extra-width-pt",
        type=float,
        default=2.0,
        help="Extra width added to each textbox",
    )
    parser.add_argument(
        "--text-box-extra-height-pt",
        type=float,
        default=8.0,
        help="Extra height added to each textbox",
    )
    parser.add_argument(
        "--enable-translate",
        action="store_true",
        help="Run the same placeholder translation layer as visual clone pipeline",
    )
    parser.add_argument(
        "--source-lang",
        default="auto",
        help="Source language used by translator placeholder",
    )
    parser.add_argument(
        "--target-lang",
        default="zh",
        help="Target language used by translator placeholder",
    )
    return parser.parse_args()


def _iter_raw_json_files(input_dir: Path) -> list[tuple[Path, str, str]]:
    matches: list[tuple[Path, str, str]] = []
    for path in sorted(input_dir.glob("*_raw.json")):
        stem = path.stem
        for engine in ENGINES:
            suffix = f"_{engine}_raw"
            if stem.endswith(suffix):
                sample_name = stem[: -len(suffix)]
                matches.append((path, sample_name, engine))
                break
    return matches


def _resolve_input_pdf(sample_name: str, search_dirs: list[Path]) -> Path:
    for root in search_dirs:
        candidate = root / f"{sample_name}.pdf"
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"Unable to locate source PDF for sample: {sample_name}")


def _load_engine_pages(payload: Any, engine: str, raw_json_path: Path | None = None) -> tuple[list[dict[str, Any]], str]:
    if engine == "paddleocr-vl":
        pages = _extract_pages(payload)
        if not pages:
            raise RuntimeError("No pages found in PaddleOCR-VL raw JSON")
        return pages, "vl"

    if engine == "mineru":
        pages = _extract_mineru_pages(payload)
        if not pages:
            pages = _extract_generic_pages(payload)
        if not pages:
            raise RuntimeError("No pages found in MinerU raw JSON")
        return pages, "generic"

    if engine == "dotsocr":
        pages = _extract_generic_pages(payload)
        if not pages:
            raise RuntimeError("No pages found in DotsOCR raw JSON")
        return pages, "generic"

    if engine == "opendataloader-hybrid-full":
        pages = _extract_opendataloader_pages(payload, base_dir=raw_json_path.parent if raw_json_path else None)
        if not pages:
            raise RuntimeError("No pages found in OpenDataLoader raw JSON")
        return pages, "opendataloader"

    raise RuntimeError(f"Unsupported engine: {engine}")


def _extract_opendataloader_pages(payload: Any, base_dir: Path | None = None) -> list[dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    kids = payload.get("kids")
    if not isinstance(kids, list):
        return []

    page_items: dict[int, list[dict[str, Any]]] = {}
    for item in kids:
        if not isinstance(item, dict):
            continue
        page_no = item.get("page number")
        if isinstance(page_no, int) and page_no >= 1:
            page_items.setdefault(page_no - 1, []).append(item)

    pages: list[dict[str, Any]] = []
    for page_index in sorted(page_items.keys()):
        raw_items = page_items[page_index]
        blocks: list[Block] = []
        images_map: dict[str, str] = {}
        order_seed = 0.0

        for item in raw_items:
            bbox = item.get("bounding box")
            if not (isinstance(bbox, list) and len(bbox) == 4):
                continue
            block_bbox = [float(v) for v in bbox]
            item_type = str(item.get("type") or "paragraph").strip().lower()

            if item_type == "image":
                image_ref = str(item.get("source") or "").strip()
                if image_ref:
                    if base_dir is not None:
                        candidate = (base_dir / image_ref).resolve()
                        images_map[image_ref] = str(candidate)
                    blocks.append(
                        Block(
                            bbox=block_bbox,
                            label="image",
                            order=order_seed,
                            text="",
                            image_ref=image_ref or None,
                        )
                    )
                order_seed += 1.0
                continue

            if item_type == "table":
                blocks.append(
                    Block(
                        bbox=block_bbox,
                        label="table",
                        order=order_seed,
                        text="",
                    )
                )
                order_seed += 1.0
                continue

            if item_type == "list":
                list_items = item.get("list items")
                if isinstance(list_items, list) and list_items:
                    for sub_idx, sub in enumerate(list_items):
                        if not isinstance(sub, dict):
                            continue
                        sub_bbox = sub.get("bounding box")
                        if not (isinstance(sub_bbox, list) and len(sub_bbox) == 4):
                            continue
                        text = str(sub.get("content") or "").strip()
                        if not text:
                            continue
                        blocks.append(
                            Block(
                                bbox=[float(v) for v in sub_bbox],
                                label="text",
                                order=order_seed + sub_idx / 100.0,
                                text=text,
                            )
                        )
                    order_seed += 1.0
                    continue

            text = str(item.get("content") or "").strip()
            if not text:
                order_seed += 1.0
                continue
            label = "paragraph_title" if item_type == "heading" else "text"
            blocks.append(
                Block(
                    bbox=block_bbox,
                    label=label,
                    order=order_seed,
                    text=text,
                )
            )
            order_seed += 1.0

        pages.append(
            {
                "blocks": blocks,
                "raw": {
                    "kids": raw_items,
                    "_coord_size_from_pdf": True,
                    "_opendataloader_bottom_left_origin": True,
                },
                "images_map": images_map,
            }
        )
    return pages


def _build_font_steps(default_font_size: float, min_font_size: float) -> list[float]:
    base_steps = [11.0, 10.0, 9.0, 8.0, 7.0, 6.0, 5.0, 4.0, 3.0, 2.5, 2.0]
    steps = [min(default_font_size, fs) for fs in base_steps if fs >= min_font_size]
    steps.append(float(default_font_size))
    filtered = sorted({round(max(min_font_size, step), 2) for step in steps if step >= min_font_size}, reverse=True)
    if not filtered:
        return [float(min_font_size)]
    return filtered


# Word 实际渲染时字符宽度比理论值偏宽，加一个保守系数避免低估换行数。
# SimSun 等中文字体的 ASCII 字符实际宽度约为字号的 0.60~0.65 倍（而非 0.52）。
_WORD_WIDTH_CORRECTION = 1.20  # 对所有字符宽度估算乘以此系数


def _estimate_wrapped_line_count(text: str, width_pt: float, font_size_pt: float, padding_pt: float) -> int:
    usable_width_pt = max(1.0, width_pt - 2.0 * max(0.0, padding_pt))
    # 用修正系数让每个字符占用更多宽度，从而估算出更多换行数，避免低估高度
    char_capacity_units = max(1.0, usable_width_pt / max(font_size_pt * _WORD_WIDTH_CORRECTION, 0.1))
    total_lines = 0
    for raw_line in (text or "").split("\n"):
        line = raw_line.rstrip("\r")
        if not line:
            total_lines += 1
            continue
        units = sum(_char_visual_unit(ch) for ch in line)
        total_lines += max(1, int((units / char_capacity_units) + 0.999))
    return max(total_lines, 1)


# Word VML textbox 每个 w:p 段落默认有段前/段后间距（约 0~2pt），
# 即使 section paragraph 设了 space_before=0，txbxContent 内部的 w:p 仍受影响。
# 用一个固定补偿值来抵消这个系统性偏差。
_WORD_PARA_SPACING_EXTRA_PT = 1.0  # 每个段落额外占用的高度补偿（pt）


def _estimate_required_height_pt(
    *,
    text: str,
    width_pt: float,
    font_size_pt: float,
    line_height: float,
    padding_pt: float,
) -> float:
    lines = _estimate_wrapped_line_count(text, width_pt, font_size_pt, padding_pt)
    # 每行高度 + 段落间距补偿 + 上下 padding
    line_h = max(0.85, font_size_pt * line_height)
    para_extra = _WORD_PARA_SPACING_EXTRA_PT * lines
    return lines * line_h + para_extra + 2.0 * max(0.0, padding_pt)


def _token_visual_units(token: str) -> float:
    return sum(_char_visual_unit(ch) for ch in token)


def _wrap_text_to_lines(
    *,
    text: str,
    width_pt: float,
    font_size_pt: float,
    padding_pt: float,
) -> list[str]:
    usable_width_pt = max(1.0, width_pt - 2.0 * max(0.0, padding_pt))
    char_capacity_units = max(1.0, usable_width_pt / max(font_size_pt * _WORD_WIDTH_CORRECTION, 0.1))
    out_lines: list[str] = []

    for raw_line in (text or "").split("\n"):
        src = raw_line.rstrip("\r")
        if not src:
            out_lines.append("")
            continue

        if " " in src:
            tokens: list[str] = []
            parts = src.split(" ")
            for i, part in enumerate(parts):
                if i > 0:
                    tokens.append(" ")
                if part:
                    tokens.append(part)
        else:
            tokens = list(src)

        current = ""
        current_units = 0.0
        for token in tokens:
            units = _token_visual_units(token)
            if current and current_units + units > char_capacity_units:
                out_lines.append(current.rstrip())
                current = token.lstrip() if token == " " else token
                current_units = _token_visual_units(current)
                continue
            current += token
            current_units += units

        if current or not out_lines:
            out_lines.append(current.rstrip())

    return out_lines or [text]


def _docx_text_fits(
    *,
    text: str,
    width_pt: float,
    height_pt: float,
    font_size_pt: float,
    line_height: float,
    padding_pt: float,
) -> bool:
    required_height = _estimate_required_height_pt(
        text=text,
        width_pt=width_pt,
        font_size_pt=font_size_pt,
        line_height=line_height,
        padding_pt=padding_pt,
    )
    return required_height <= max(1.0, height_pt)


def _choose_font_size_pt(
    *,
    text: str,
    label: str,
    width_pt: float,
    height_pt: float,
    default_font_size: float,
    min_font_size: float,
    line_height: float,
    padding_pt: float,
) -> float:
    target_default = float(default_font_size)
    if label == "paragraph_title":
        target_default = min(default_font_size + 1.0, 12.0)
    elif label == "doc_title":
        target_default = min(default_font_size + 2.0, 13.0)

    for font_size in _build_font_steps(target_default, min_font_size):
        if _docx_text_fits(
            text=text,
            width_pt=width_pt,
            height_pt=height_pt,
            font_size_pt=font_size,
            line_height=line_height,
            padding_pt=padding_pt,
        ):
            return font_size

    return float(min_font_size)


def _shrink_font_size_to_fit_height(
    *,
    text: str,
    width_pt: float,
    available_height_pt: float,
    current_font_size: float,
    min_font_size: float,
    line_height: float,
    padding_pt: float,
) -> float:
    if available_height_pt <= 0:
        return max(min_font_size, current_font_size)
    if _docx_text_fits(
        text=text,
        width_pt=width_pt,
        height_pt=available_height_pt,
        font_size_pt=current_font_size,
        line_height=line_height,
        padding_pt=padding_pt,
    ):
        return current_font_size
    for font_size in _build_font_steps(current_font_size, min_font_size):
        if _docx_text_fits(
            text=text,
            width_pt=width_pt,
            height_pt=available_height_pt,
            font_size_pt=font_size,
            line_height=line_height,
            padding_pt=padding_pt,
        ):
            return font_size
    return float(min_font_size)


def _adaptive_profile(rect: fitz.Rect, engine: str = "") -> dict[str, float | bool]:
    short_edge = min(rect.width, rect.height)
    page_area = rect.width * rect.height
    is_dense_small_page = short_edge < 520 or page_area < 180000

    # mineru 倾向于把文本拆成更细的块（同样页面块数约是 dotsocr 的 1.5x），
    # 相邻块间距更小，extra_height 需要收紧以减少重叠。
    if engine == "mineru":
        extra_w = 0.4 if is_dense_small_page else 1.2
        extra_h = 0.8 if is_dense_small_page else 2.0
        text_box_padding = 1.0 if is_dense_small_page else 2.0
        min_font = 2.0 if is_dense_small_page else 2.5
        line_height = 0.9 if is_dense_small_page else 1.0
    elif engine == "dotsocr":
        # DotsOCR 长段较多：增高文本框并提高最小字号，减少“字太小+段过长出界”。
        extra_w = 1.0 if is_dense_small_page else 2.4
        extra_h = 3.0 if is_dense_small_page else 8.0
        text_box_padding = 0.6 if is_dense_small_page else 1.2
        min_font = 5.8 if is_dense_small_page else 6.4
        line_height = 0.86 if is_dense_small_page else 0.92
    elif engine == "paddleocr-vl":
        # PaddleOCR-VL 常见块高估不足，适当增高并抬高最小字号，避免页底小字堆叠。
        extra_w = 1.2 if is_dense_small_page else 3.0
        extra_h = 2.4 if is_dense_small_page else 7.0
        text_box_padding = 0.8 if is_dense_small_page else 1.4
        min_font = 5.2 if is_dense_small_page else 5.8
        line_height = 0.88 if is_dense_small_page else 0.95
    else:
        extra_w = 0.8 if is_dense_small_page else 2.0
        extra_h = 1.5 if is_dense_small_page else 5.0
        text_box_padding = 1.0 if is_dense_small_page else 2.0
        min_font = 2.0 if is_dense_small_page else 2.5
        line_height = 0.9 if is_dense_small_page else 1.0

    return {
        "textBoxPadding": text_box_padding,
        "minFontSize": min_font,
        "lineHeight": line_height,
        "receiptMode": is_dense_small_page,
        "textBoxExtraWidthPt": extra_w,
        "textBoxExtraHeightPt": extra_h,
    }


def _unwrap_raw_meta(raw: Any) -> dict[str, Any]:
    current = raw
    for _ in range(3):
        if not isinstance(current, dict):
            return {}
        nested = current.get("raw")
        if not isinstance(nested, dict):
            return current
        current = nested
    return current if isinstance(current, dict) else {}


def _resolve_overlap_top(
    *,
    left_pt: float,
    width_pt: float,
    top_pt: float,
    height_pt: float,
    placed: list[tuple[float, float, float, float]],
    page_height: float,
    margin: float = 1.5,
    max_shift: float | None = None,
) -> float:
    """如果当前文本框与已放置的框有垂直重叠，向下推移到不重叠的位置。

    只在与当前框有实际横向交叠的情况下做 y 轴推移，避免把不同列的内容也当成冲突。
    """
    x1_cur = left_pt
    x2_cur = left_pt + max(width_pt, 1.0)
    new_top = top_pt
    # 默认允许在页内尽可能下移，优先满足“不重叠”。
    allowed_shift = max_shift if max_shift is not None else page_height
    sorted_placed = sorted(placed, key=lambda box: (box[1], box[0]))
    for px1, py1, px2, py2 in sorted_placed:
        # 检查水平方向是否有实质重叠（超过各自宽度的 20%）
        overlap_x = min(x2_cur, px2) - max(x1_cur, px1)
        if overlap_x <= 0:
            continue
        min_width = max(1.0, min(x2_cur - x1_cur, px2 - px1))
        if (overlap_x / min_width) < 0.08:
            continue
        # 检查垂直方向是否重叠
        cur_y2 = new_top + height_pt
        if new_top < py2 - margin and cur_y2 > py1 + margin:
            # 向下推到已放置框的底部 + margin
            candidate = py2 + margin
            if candidate - top_pt <= allowed_shift:
                new_top = max(new_top, candidate)
    # 不超出页面
    new_top = min(new_top, max(0.0, page_height - height_pt))
    return new_top


def _build_tiny_box(
    *,
    x1: float,
    y1: float,
    width_pt: float,
    height_pt: float,
    page_rect: fitz.Rect,
    text: str,
    font_size_pt: float,
) -> tuple[float, float, float, float]:
    draw_width = max(width_pt, min(page_rect.width - x1, max(8.0, len(text) * font_size_pt * 0.56)))
    draw_height = max(height_pt, font_size_pt * 1.15 + 1.0)
    draw_x = min(max(0.5, x1), max(0.5, page_rect.width - 2.0))
    draw_y = min(max(0.5, y1), max(0.5, page_rect.height - draw_height))
    if draw_x + draw_width > page_rect.width:
        draw_width = max(2.0, page_rect.width - draw_x)
    if draw_y + draw_height > page_rect.height:
        draw_height = max(2.0, page_rect.height - draw_y)
    return draw_x, draw_y, draw_width, draw_height


def _set_section_size(doc: Document, page_index: int, page_width_pt: float, page_height_pt: float):
    if page_index == 0:
        section = doc.sections[0]
    else:
        section = doc.add_section()
    section.page_width = Pt(page_width_pt)
    section.page_height = Pt(page_height_pt)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def _vl_coord_size(page_item: dict[str, Any], rect: fitz.Rect) -> tuple[float, float]:
    # 坐标尺度优先取“图像尺寸 + 块坐标上界”的较大值，避免尺度偏小导致整体右移。
    candidates: list[tuple[float, float]] = []
    image_size = _get_input_image_size(page_item)
    if image_size and image_size[0] > 0 and image_size[1] > 0:
        candidates.append((float(image_size[0]), float(image_size[1])))

    pruned = page_item.get("prunedResult", {}) if isinstance(page_item, dict) else {}
    max_x = 0.0
    max_y = 0.0
    if isinstance(pruned, dict):
        parsing = pruned.get("parsing_res_list")
        if isinstance(parsing, list):
            for item in parsing:
                if not isinstance(item, dict):
                    continue
                bbox = item.get("block_bbox")
                if isinstance(bbox, list) and len(bbox) == 4:
                    try:
                        max_x = max(max_x, float(bbox[2]))
                        max_y = max(max_y, float(bbox[3]))
                    except Exception:
                        pass
        layout_boxes = pruned.get("layout_det_res", {}).get("boxes", [])
        if isinstance(layout_boxes, list):
            for item in layout_boxes:
                if not isinstance(item, dict):
                    continue
                coord = item.get("coordinate")
                if isinstance(coord, list) and len(coord) == 4:
                    try:
                        max_x = max(max_x, float(coord[2]))
                        max_y = max(max_y, float(coord[3]))
                    except Exception:
                        pass
    if max_x > 0 and max_y > 0:
        candidates.append((max_x, max_y))

    if not candidates:
        return rect.width, rect.height
    coord_w = max(rect.width, max(w for w, _ in candidates))
    coord_h = max(rect.height, max(h for _, h in candidates))
    return coord_w, coord_h


def _page_blocks(
    *,
    page_item: dict[str, Any],
    provider_kind: str,
    geometry_source: str,
    rect: fitz.Rect,
) -> tuple[list[Any], str, dict[str, Any], float, float]:
    if provider_kind == "vl":
        coord_size = _vl_coord_size(page_item, rect)
        coord_w, coord_h = coord_size
        images_map = page_item.get("markdown", {}).get("images", {})
        if not isinstance(images_map, dict):
            images_map = {}
    elif provider_kind == "opendataloader":
        coord_w, coord_h = _generic_page_coord_size(page_item, rect)
        images_map = page_item.get("images_map", {})
        if not isinstance(images_map, dict):
            images_map = {}
    else:
        coord_w, coord_h = _generic_page_coord_size(page_item, rect)
        images_map = {}

    sx = rect.width / max(coord_w, 1.0)
    sy = rect.height / max(coord_h, 1.0)
    selected_geometry = geometry_source

    if provider_kind == "vl":
        if geometry_source == "layout-hybrid":
            blocks = _build_hybrid_blocks(page_item, keep_unmatched_chunks=True)
        elif geometry_source == "auto":
            blocks, selected_geometry, _ = _choose_blocks_auto(page_item, rect=rect, sx=sx, sy=sy)
        else:
            blocks = _build_parsing_blocks(page_item)
    else:
        blocks = page_item.get("blocks", [])
        if geometry_source != "parsing":
            selected_geometry = "parsing"

    if provider_kind == "opendataloader" and _opendataloader_uses_bottom_left_origin(page_item):
        blocks = [_clone_block_with_bbox(block, _flip_bbox_vertical(block.bbox, coord_h)) for block in blocks]

    return blocks, selected_geometry, images_map, sx, sy


def _generic_page_coord_size(page_item: dict[str, Any], rect: fitz.Rect) -> tuple[float, float]:
    raw = _unwrap_raw_meta(page_item.get("raw"))
    if raw:
        if raw.get("_coord_size_from_pdf"):
            return rect.width, rect.height
        page_size = raw.get("page_size")
        if isinstance(page_size, list) and len(page_size) == 2:
            try:
                w = float(page_size[0])
                h = float(page_size[1])
                if w > 0 and h > 0:
                    return w, h
            except Exception:
                pass

        iw = raw.get("input_width")
        ih = raw.get("input_height")
        if isinstance(iw, (int, float)) and isinstance(ih, (int, float)) and iw > 0 and ih > 0:
            return float(iw), float(ih)

        layout_image_path = raw.get("layout_image_path")
        if isinstance(layout_image_path, str) and layout_image_path.strip():
            p = Path(layout_image_path)
            if p.exists():
                try:
                    pix = fitz.Pixmap(str(p))
                    if pix.width > 0 and pix.height > 0:
                        return float(pix.width), float(pix.height)
                except Exception:
                    pass

    generic_blocks = page_item.get("blocks", [])
    return _choose_generic_coord_size(generic_blocks, rect)


def _is_visual_only_label(label: str) -> bool:
    return str(label).strip().lower() in VISUAL_ONLY_LABELS


def _render_crop_bytes(page: fitz.Page, rect: fitz.Rect) -> bytes | None:
    clip = fitz.Rect(
        max(0.0, rect.x0),
        max(0.0, rect.y0),
        min(page.rect.width, rect.x1),
        min(page.rect.height, rect.y1),
    )
    if clip.width <= 1.0 or clip.height <= 1.0:
        return None
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), clip=clip, alpha=False)
        return pix.tobytes("png")
    except Exception:
        return None


def _expand_text_box_height(
    *,
    text: str,
    width_pt: float,
    base_height_pt: float,
    font_size_pt: float,
    line_height: float,
    padding_pt: float,
    page_height_pt: float,
    top_pt: float,
) -> float:
    required_height = _estimate_required_height_pt(
        text=text,
        width_pt=width_pt,
        font_size_pt=font_size_pt,
        line_height=line_height,
        padding_pt=padding_pt,
    )
    safety_height = required_height + max(1.5, font_size_pt * 0.2)
    target_height = max(base_height_pt, safety_height)
    max_draw_height = max(2.0, page_height_pt - top_pt)
    return min(target_height, max_draw_height)


def _vertical_overlap_ratio(y1: float, y2: float, oy1: float, oy2: float) -> float:
    overlap = min(y2, oy2) - max(y1, oy1)
    if overlap <= 0:
        return 0.0
    min_height = max(1.0, min(y2 - y1, oy2 - oy1))
    return overlap / min_height


def _is_textbox_label(label: str) -> bool:
    normalized = str(label).strip().lower()
    return normalized not in VISUAL_ONLY_LABELS


def _is_title_like_label(label: str) -> bool:
    normalized = str(label).strip().lower()
    return normalized in {"doc_title", "paragraph_title", "title", "heading", "header"}


def _detect_column_regions(
    mapped_blocks: list[dict[str, Any]],
    page_width_pt: float,
) -> list[dict[str, float]]:
    text_blocks = [b for b in mapped_blocks if _is_textbox_label(str(b.get("label", "")))]
    if len(text_blocks) < 2:
        return []

    regions: list[dict[str, float]] = []
    min_column_width = page_width_pt * 0.18
    min_gap = page_width_pt * 0.04

    for left in text_blocks:
        lx1 = float(left["x1"])
        lx2 = float(left["x2"])
        lw = lx2 - lx1
        if lw < min_column_width or lx1 > page_width_pt * 0.42:
            continue
        for right in text_blocks:
            if left is right:
                continue
            rx1 = float(right["x1"])
            rx2 = float(right["x2"])
            rw = rx2 - rx1
            if rw < min_column_width or rx1 < page_width_pt * 0.45:
                continue
            overlap_ratio = _vertical_overlap_ratio(
                float(left["y1"]),
                float(left["y2"]),
                float(right["y1"]),
                float(right["y2"]),
            )
            if overlap_ratio < 0.55:
                continue
            gap = rx1 - lx2
            if gap < min_gap:
                continue
            band_top = min(float(left["y1"]), float(right["y1"]))
            band_bottom = max(float(left["y2"]), float(right["y2"]))
            split_x = (lx2 + rx1) / 2.0
            regions.append(
                {
                    "top": band_top,
                    "bottom": band_bottom,
                    "left_x1": min(lx1, float(right["x1"])),
                    "left_x2": split_x,
                    "right_x1": split_x,
                    "right_x2": max(float(left["x2"]), rx2),
                }
            )

    if not regions:
        return []

    regions.sort(key=lambda item: (item["top"], item["bottom"]))
    merged: list[dict[str, float]] = []
    for region in regions:
        if not merged:
            merged.append(region)
            continue
        prev = merged[-1]
        same_split = abs((prev["left_x2"] - prev["right_x1"]) - (region["left_x2"] - region["right_x1"])) < 24.0
        vertical_near = region["top"] <= prev["bottom"] + 18.0
        if same_split and vertical_near:
            prev["top"] = min(prev["top"], region["top"])
            prev["bottom"] = max(prev["bottom"], region["bottom"])
            prev["left_x1"] = min(prev["left_x1"], region["left_x1"])
            prev["left_x2"] = min(prev["left_x2"], region["left_x2"])
            prev["right_x1"] = max(prev["right_x1"], region["right_x1"])
            prev["right_x2"] = max(prev["right_x2"], region["right_x2"])
        else:
            merged.append(region)
    return merged


def _column_limited_width(
    *,
    current: dict[str, Any],
    desired_width_pt: float,
    column_regions: list[dict[str, float]],
    page_width_pt: float,
    gutter_pt: float = 6.0,
) -> float:
    x1 = float(current["x1"])
    x2 = float(current["x2"])
    y1 = float(current["y1"])
    y2 = float(current["y2"])
    center_x = (x1 + x2) / 2.0
    limited_right = page_width_pt - 0.5

    for region in column_regions:
        overlap_ratio = _vertical_overlap_ratio(y1, y2, float(region["top"]), float(region["bottom"]))
        if overlap_ratio < 0.25:
            continue
        split_x = (float(region["left_x2"]) + float(region["right_x1"])) / 2.0
        if center_x <= split_x:
            limited_right = min(limited_right, float(region["left_x2"]) - gutter_pt)
        else:
            left_bound = float(region["right_x1"]) + gutter_pt
            if x1 < left_bound:
                x1 = left_bound
            limited_right = min(limited_right, float(region["right_x2"]) - gutter_pt)

    return max(2.0, min(desired_width_pt, limited_right - float(current["x1"])))


def _apply_column_bounds(
    *,
    current: dict[str, Any],
    desired_left_pt: float,
    desired_width_pt: float,
    column_regions: list[dict[str, float]],
    page_width_pt: float,
    gutter_pt: float = 6.0,
) -> tuple[float, float, str]:
    x1 = float(current["x1"])
    x2 = float(current["x2"])
    y1 = float(current["y1"])
    y2 = float(current["y2"])
    center_x = (x1 + x2) / 2.0

    best_bucket = "full"
    left_bound = max(0.0, desired_left_pt)
    right_bound = page_width_pt - 0.5
    best_overlap = 0.0

    for idx, region in enumerate(column_regions):
        overlap_ratio = _vertical_overlap_ratio(y1, y2, float(region["top"]), float(region["bottom"]))
        if overlap_ratio < 0.25 or overlap_ratio < best_overlap:
            continue
        split_x = (float(region["left_x2"]) + float(region["right_x1"])) / 2.0
        best_overlap = overlap_ratio
        if center_x <= split_x:
            best_bucket = f"col_{idx}_left"
            left_bound = max(left_bound, float(region["left_x1"]))
            right_bound = min(right_bound, float(region["left_x2"]) - gutter_pt)
        else:
            best_bucket = f"col_{idx}_right"
            left_bound = max(left_bound, float(region["right_x1"]) + gutter_pt)
            right_bound = min(right_bound, float(region["right_x2"]) - gutter_pt)

    if right_bound <= left_bound + 2.0:
        right_bound = min(page_width_pt - 0.5, left_bound + max(2.0, desired_width_pt))

    draw_left = min(max(desired_left_pt, left_bound), max(left_bound, right_bound - 2.0))
    draw_width = max(2.0, min(desired_width_pt, right_bound - draw_left))
    return draw_left, draw_width, best_bucket


def _limit_text_box_width(
    *,
    block_index: int,
    mapped_blocks: list[dict[str, Any]],
    column_regions: list[dict[str, float]],
    page_width_pt: float,
    desired_width_pt: float,
    gutter_pt: float = 6.0,
) -> float:
    current = mapped_blocks[block_index]
    x1 = float(current["x1"])
    y1 = float(current["y1"])
    y2 = float(current["y2"])

    page_right = max(x1 + 2.0, page_width_pt - 0.5)
    max_right = page_right

    column_limited = _column_limited_width(
        current=current,
        desired_width_pt=desired_width_pt,
        column_regions=column_regions,
        page_width_pt=page_width_pt,
        gutter_pt=gutter_pt,
    )
    max_right = min(max_right, x1 + column_limited)

    for idx, other in enumerate(mapped_blocks):
        if idx == block_index:
            continue
        ox1 = float(other["x1"])
        if ox1 <= x1 + 8.0:
            continue
        overlap_ratio = _vertical_overlap_ratio(y1, y2, float(other["y1"]), float(other["y2"]))
        if overlap_ratio < 0.18:
            continue
        candidate_right = ox1 - gutter_pt
        if candidate_right <= x1 + 2.0:
            continue
        max_right = min(max_right, candidate_right)

    return max(2.0, min(desired_width_pt, max_right - x1))


def _should_render_text_as_image(
    *,
    label: str,
    text: str,
    width_pt: float,
    height_pt: float,
) -> bool:
    if not text.strip():
        return False
    if not _is_textbox_label(label):
        return False
    compact_text = text.replace("\n", "").strip()
    if len(compact_text) > 40:
        return False
    if width_pt <= 0 or height_pt <= 0:
        return False
    aspect_ratio = height_pt / max(width_pt, 1.0)
    has_spaces = any(ch.isspace() for ch in compact_text)
    return aspect_ratio >= 1.7 and (not has_spaces or len(compact_text) <= 16)


def _engine_textbox_safe_width(
    *,
    engine: str,
    left_pt: float,
    width_pt: float,
    page_width_pt: float,
    bucket: str,
) -> tuple[float, float]:
    right_margin = 6.0
    shrink_factor = 1.0

    if engine == "dotsocr":
        shrink_factor = 0.985
        right_margin = 4.0
    elif engine == "paddleocr-vl":
        shrink_factor = 0.99
        right_margin = 4.0
    elif engine == "mineru":
        shrink_factor = 0.96
        right_margin = 6.0

    if bucket.endswith("_right"):
        shrink_factor = min(shrink_factor, 0.96 if engine == "dotsocr" else 0.97)
        right_margin = max(right_margin, 6.0)
    elif bucket.endswith("_left"):
        shrink_factor = min(shrink_factor, 0.97 if engine == "dotsocr" else 0.98)
        right_margin = max(right_margin, 5.0)

    max_width = max(2.0, page_width_pt - left_pt - right_margin)
    safe_width = min(max_width, max(2.0, width_pt * shrink_factor))
    return safe_width, right_margin


def _bbox_close(a: tuple[float, float, float, float], b: tuple[float, float, float, float], tol: float = 3.0) -> bool:
    return all(abs(x - y) <= tol for x, y in zip(a, b))


def _bbox_overlap(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> bool:
    overlap_x = min(a[2], b[2]) - max(a[0], b[0])
    overlap_y = min(a[3], b[3]) - max(a[1], b[1])
    return overlap_x > 0 and overlap_y > 0


def _bbox_intersection_area(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    overlap_x = min(a[2], b[2]) - max(a[0], b[0])
    overlap_y = min(a[3], b[3]) - max(a[1], b[1])
    if overlap_x <= 0 or overlap_y <= 0:
        return 0.0
    return overlap_x * overlap_y


def _bbox_overlap_ratio(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    inter = _bbox_intersection_area(a, b)
    if inter <= 0:
        return 0.0
    area_a = max(1.0, (a[2] - a[0]) * (a[3] - a[1]))
    area_b = max(1.0, (b[2] - b[0]) * (b[3] - b[1]))
    return inter / min(area_a, area_b)


def _is_compact_heading_text(text: str) -> bool:
    normalized = " ".join((text or "").split()).strip()
    if not normalized:
        return False
    if len(normalized) > 36:
        return False
    # 典型编号标题："1. Headings" / "2) Methods"
    if re.match(r"^\d+[\.|\)]\s+\S+", normalized):
        return True
    # 简短标题短语，避免被误拆成多行导致后续遮挡/截断
    words = normalized.split(" ")
    if 1 <= len(words) <= 4 and all(len(w) <= 16 for w in words):
        alpha_words = [w for w in words if re.search(r"[A-Za-z]", w)]
        if alpha_words and sum(1 for w in alpha_words if w[:1].isupper()) >= 1:
            return True
    return False


def _looks_like_bullet_list_text(text: str) -> bool:
    lines = [ln.strip() for ln in (text or "").split("\n") if ln.strip()]
    if len(lines) < 2:
        return False
    bullet_prefix = re.compile(r"^[\u2022\u25A0\u25A1\u25CF\u25CB\-\*]\s+\S+")
    bullet_count = sum(1 for ln in lines if bullet_prefix.match(ln))
    return bullet_count >= 2


def _collect_mineru_table_html_candidates(obj: Any, candidates: list[dict[str, Any]]) -> None:
    if isinstance(obj, dict):
        obj_type = str(obj.get("type") or "").strip().lower()
        bbox = obj.get("bbox")
        if obj_type == "table" and isinstance(bbox, list) and len(bbox) == 4:
            html_hits: list[str] = []

            def walk_html(node: Any) -> None:
                if isinstance(node, dict):
                    html = node.get("html")
                    if isinstance(html, str) and "<table" in html.lower():
                        html_hits.append(html)
                    for value in node.values():
                        walk_html(value)
                elif isinstance(node, list):
                    for item in node:
                        walk_html(item)

            walk_html(obj)
            for html in html_hits:
                candidates.append({"bbox": tuple(float(v) for v in bbox), "html": html})
        for value in obj.values():
            _collect_mineru_table_html_candidates(value, candidates)
    elif isinstance(obj, list):
        for item in obj:
            _collect_mineru_table_html_candidates(item, candidates)


def _page_table_html_candidates(
    page_item: dict[str, Any],
    provider_kind: str,
    coord_h: float | None = None,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    if provider_kind == "opendataloader":
        raw = page_item.get("raw", {})
        kids = raw.get("kids") if isinstance(raw, dict) else None
        if isinstance(kids, list):
            for item in kids:
                if not isinstance(item, dict):
                    continue
                if str(item.get("type") or "").strip().lower() != "table":
                    continue
                bbox = item.get("bounding box")
                if not (isinstance(bbox, list) and len(bbox) == 4):
                    continue
                rows = _table_rows_from_opendataloader_item(item)
                if rows:
                    candidate_bbox = [float(v) for v in bbox]
                    if coord_h is not None and _opendataloader_uses_bottom_left_origin(page_item):
                        candidate_bbox = _flip_bbox_vertical(candidate_bbox, coord_h)
                    candidates.append({"bbox": tuple(candidate_bbox), "rows": rows})
        return candidates

    if provider_kind == "vl":
        pruned = page_item.get("prunedResult", {})
        parsing = pruned.get("parsing_res_list", [])
        if isinstance(parsing, list):
            for item in parsing:
                if not isinstance(item, dict):
                    continue
                if str(item.get("block_label", "")).strip().lower() != "table":
                    continue
                html = str(item.get("block_content", "") or "")
                bbox = item.get("block_bbox")
                if "<table" in html.lower() and isinstance(bbox, list) and len(bbox) == 4:
                    candidates.append({"bbox": tuple(float(v) for v in bbox), "html": html})
        return candidates

    raw = page_item.get("raw")
    if isinstance(raw, dict):
        result = raw.get("result")
        if isinstance(result, list):
            for item in result:
                if not isinstance(item, dict):
                    continue
                label = str(item.get("category_type") or item.get("category") or item.get("label") or "").strip().lower()
                if label != "table":
                    continue
                html = str(item.get("text", "") or item.get("content", "") or "")
                bbox = item.get("bbox") or item.get("bbox_2d")
                if "<table" in html.lower() and isinstance(bbox, list) and len(bbox) == 4:
                    candidates.append({"bbox": tuple(float(v) for v in bbox), "html": html})
        _collect_mineru_table_html_candidates(raw, candidates)
    return candidates


def _match_table_html(
    *,
    block: Any,
    table_candidates: list[dict[str, Any]],
) -> dict[str, Any] | None:
    bbox = tuple(float(v) for v in block.bbox)
    for idx, cand in enumerate(table_candidates):
        if _bbox_close(bbox, cand["bbox"]):
            return table_candidates.pop(idx)
    for idx, cand in enumerate(table_candidates):
        cb = cand["bbox"]
        overlap_x = min(bbox[2], cb[2]) - max(bbox[0], cb[0])
        overlap_y = min(bbox[3], cb[3]) - max(bbox[1], cb[1])
        if overlap_x > 0 and overlap_y > 0:
            return table_candidates.pop(idx)
    return None


def _render_word_table(doc: Document, table_rows: list[list[dict[str, Any]]], font_name: str, font_size_pt: float) -> bool:
    if not table_rows:
        return False
    max_cols = max(sum(max(1, int(cell.get("colspan", 1))) for cell in row) for row in table_rows)
    if max_cols <= 0:
        return False

    table = doc.add_table(rows=len(table_rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    occupied: dict[tuple[int, int], bool] = {}
    for r_idx, row in enumerate(table_rows):
        c_idx = 0
        for cell in row:
            while occupied.get((r_idx, c_idx)):
                c_idx += 1
            base = table.cell(r_idx, c_idx)
            text = str(cell.get("text", "") or "")
            base.text = text
            for paragraph in base.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)

            rowspan = max(1, int(cell.get("rowspan", 1)))
            colspan = max(1, int(cell.get("colspan", 1)))

            end_row = min(len(table_rows) - 1, r_idx + rowspan - 1)
            end_col = min(max_cols - 1, c_idx + colspan - 1)
            if end_col > c_idx:
                base = base.merge(table.cell(r_idx, end_col))
            if end_row > r_idx:
                base = base.merge(table.cell(end_row, end_col))

            for rr in range(r_idx, end_row + 1):
                for cc in range(c_idx, end_col + 1):
                    occupied[(rr, cc)] = True
            c_idx = end_col + 1
    return True


def _add_vertical_spacer(doc: Document, height_pt: float) -> None:
    spacer_height = max(0.0, float(height_pt))
    if spacer_height < 1.0:
        return
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(spacer_height)
    run = paragraph.add_run(" ")
    run.font.size = Pt(1)


def _should_split_into_line_boxes(
    *,
    engine: str,
    label: str,
    text: str,
    width_pt: float,
    bucket: str,
) -> bool:
    if engine not in {"dotsocr", "paddleocr-vl", "mineru"}:
        return False
    if not _is_textbox_label(label):
        return False
    compact = text.replace("\n", " ").strip()
    if engine == "dotsocr":
        min_chars, min_width = 8, 80.0
    elif engine == "paddleocr-vl":
        min_chars, min_width = 16, 120.0
    else:  # mineru
        min_chars, min_width = 14, 100.0
    if len(compact) < min_chars:
        return False
    if bucket != "full":
        return True
    return width_pt >= min_width


def _looks_like_link_text(text: str) -> bool:
    normalized = " ".join((text or "").split()).strip()
    if not normalized:
        return False
    lowered = normalized.lower()
    if re.search(r"(https?://|www\.|[\w.+-]+@[\w-]+\.[\w.-]+)", lowered):
        return True
    if normalized in {"Top of this Page", "Sample Document", "Sample Document (docx)"}:
        return True
    if "docx" in lowered or "pdf" in lowered:
        return True
    return False


def _extract_markdown_heading(text: str) -> tuple[str, int]:
    src = (text or "").strip()
    if not src:
        return text, 0
    first, sep, rest = src.partition("\n")
    m = re.match(r"^\s*(#{1,6})\s+(.+?)\s*$", first)
    if not m:
        return text, 0
    level = len(m.group(1))
    heading_text = m.group(2).strip()
    if rest.strip():
        return f"{heading_text}\n{rest.strip()}", level
    return heading_text, level


def _export_single_docx(
    *,
    payload: Any,
    engine: str,
    raw_json_path: Path | None,
    input_pdf: Path,
    output_docx: Path,
    geometry_source: str,
    with_page_background: bool,
    font_name: str,
    font_size: float,
    min_font_size: float,
    text_box_extra_width_pt: float,
    text_box_extra_height_pt: float,
    translator: Translator,
) -> dict[str, Any]:
    pages, provider_kind = _load_engine_pages(payload, engine, raw_json_path=raw_json_path)
    src_doc = fitz.open(str(input_pdf))
    doc = Document()
    page_count = min(len(pages), len(src_doc))
    rendered_pages = 0
    rendered_blocks = 0

    for page_index in range(page_count):
        page_item = pages[page_index]
        src_page = src_doc[page_index]
        rect = src_page.rect
        # 传入 engine 让 adaptive_profile 按引擎调整 extra_height 参数
        adaptive_profile = _adaptive_profile(rect, engine=engine)
        blocks, selected_geometry, images_map, sx, sy = _page_blocks(
            page_item=page_item,
            provider_kind=provider_kind,
            geometry_source=geometry_source,
            rect=rect,
        )

        paragraph = _set_section_size(doc, page_index, rect.width, rect.height)
        if with_page_background:
            bg_bytes = src_page.get_pixmap(alpha=False).tobytes("png")
            add_image_box(
                paragraph,
                image_bytes=bg_bytes,
                left_pt=0.0,
                top_pt=0.0,
                width_pt=rect.width,
                height_pt=rect.height,
                z_index=0,
            )

        image_z_index = 10
        text_z_index = 10000
        # 按区域记录已放置的文本框，双栏页左右列分别避让。
        placed_boxes_map: dict[str, list[tuple[float, float, float, float]]] = {"full": []}
        # 全页已放置文本框：用于跨列/跨桶碰撞兜底，避免漏检重叠。
        placed_boxes_all: list[tuple[float, float, float, float]] = []
        mapped_blocks: list[dict[str, Any]] = []
        pending_tables: list[dict[str, Any]] = []
        pending_table_bboxes: list[tuple[float, float, float, float]] = []
        rendered_image_like_bboxes: list[tuple[float, float, float, float]] = []
        occupied_zones: list[tuple[float, float, float, float]] = []
        table_coord_h = _generic_page_coord_size(page_item, rect)[1] if provider_kind == "opendataloader" else None
        table_candidates = _page_table_html_candidates(page_item, provider_kind, coord_h=table_coord_h)

        for block in blocks:
            mx1 = max(0.0, float(block.bbox[0]) * sx)
            my1 = max(0.0, float(block.bbox[1]) * sy)
            mx2 = min(rect.width, float(block.bbox[2]) * sx)
            my2 = min(rect.height, float(block.bbox[3]) * sy)
            mapped_blocks.append(
                {
                    "block": block,
                    "x1": mx1,
                    "y1": my1,
                    "x2": mx2,
                    "y2": my2,
                    "label": str(block.label),
                }
            )

        # PaddleOCR-VL 在部分页面存在系统性右偏：做一次页级左移校正。
        if engine == "paddleocr-vl":
            text_boxes = [m for m in mapped_blocks if _is_textbox_label(str(m.get("label", "")))]
            if text_boxes:
                min_x = min(float(m["x1"]) for m in text_boxes)
                target_left = rect.width * 0.07
                if min_x > target_left + 20.0:
                    shift = min_x - target_left
                    for m in mapped_blocks:
                        m["x1"] = max(0.0, float(m["x1"]) - shift)
                        m["x2"] = min(rect.width, max(float(m["x1"]) + 2.0, float(m["x2"]) - shift))

        text_source_bboxes: list[tuple[float, float, float, float]] = []
        for mapped in mapped_blocks:
            if not _is_textbox_label(str(mapped.get("label", ""))):
                continue
            block = mapped.get("block")
            block_text = ""
            if block is not None:
                block_text = str(getattr(block, "text", "") or "").strip()
            if not block_text:
                continue
            text_source_bboxes.append((float(mapped["x1"]), float(mapped["y1"]), float(mapped["x2"]), float(mapped["y2"])))

        column_regions = _detect_column_regions(mapped_blocks, rect.width)
        if engine == "mineru":
            # MinerU 行级块在本项目中不稳定触发伪双栏，直接禁用列裁剪。
            column_regions = []

        # 先整页收集表格区域，避免“文本在前、表格在后”时漏掉文本-表格碰撞避让。
        table_rows_by_index: dict[int, list[list[str]]] = {}
        for block_index, mapped in enumerate(mapped_blocks):
            block = mapped["block"]
            label = str(mapped["label"])
            if label.strip().lower() != "table":
                continue
            table_candidate = _match_table_html(block=block, table_candidates=table_candidates)
            if not table_candidate:
                continue
            parsed_table: list[list[str]] = []
            if isinstance(table_candidate.get("rows"), list):
                parsed_table = table_candidate["rows"]
            else:
                parsed_table = _parse_html_table(str(table_candidate.get("html") or ""))
            if not parsed_table:
                continue
            table_rows_by_index[block_index] = parsed_table
            table_margin = 1.2
            pending_table_bboxes.append(
                (
                    max(0.0, float(mapped["x1"]) - table_margin),
                    max(0.0, float(mapped["y1"]) - table_margin),
                    min(rect.width, float(mapped["x2"]) + table_margin),
                    min(rect.height, float(mapped["y2"]) + table_margin),
                )
            )

        # 统一碰撞集合：先放入表格禁入区，后续文本框与图片都加入该集合做完整检测。
        occupied_zones.extend(pending_table_bboxes)

        for block_index, mapped in enumerate(mapped_blocks):
            block = mapped["block"]
            x1 = float(mapped["x1"])
            y1 = float(mapped["y1"])
            x2 = float(mapped["x2"])
            y2 = float(mapped["y2"])
            label = str(mapped["label"])
            width_pt = max(
                (x2 - x1) + min(text_box_extra_width_pt, float(adaptive_profile["textBoxExtraWidthPt"])),
                2.0,
            )
            base_height_pt = max(
                (y2 - y1) + min(text_box_extra_height_pt, float(adaptive_profile["textBoxExtraHeightPt"])),
                2.0,
            )
            tiny_box = (x2 - x1) < 1.0 or (y2 - y1) < 1.0

            if block_index in table_rows_by_index:
                pending_tables.append(
                    {
                        "rows": table_rows_by_index[block_index],
                        "top": y1,
                        "height": max(y2 - y1, 2.0),
                    }
                )
                continue

            rendered_visual = False
            block_bbox = (x1, y1, x2, y2)
            normalized_label = label.strip().lower()
            if block.image_ref:
                image_bytes = _resolve_image_bytes(block.image_ref, images_map)
                if image_bytes:
                    img_width_pt = max(x2 - x1, 2.0)
                    img_height_pt = max(y2 - y1, 2.0)
                    img_top_pt = _resolve_overlap_top(
                        left_pt=x1,
                        width_pt=img_width_pt,
                        top_pt=y1,
                        height_pt=img_height_pt,
                        placed=occupied_zones,
                        page_height=rect.height,
                        margin=0.4,
                        max_shift=None,
                    )
                    add_image_box(
                        paragraph,
                        image_bytes=image_bytes,
                        left_pt=x1,
                        top_pt=img_top_pt,
                        width_pt=img_width_pt,
                        height_pt=img_height_pt,
                        z_index=image_z_index,
                    )
                    occupied_zones.append((x1, img_top_pt, x1 + img_width_pt, img_top_pt + img_height_pt))
                    image_z_index += 1
                    rendered_visual = True
                    if normalized_label in IMAGE_LIKE_LABELS:
                        rendered_image_like_bboxes.append((x1, img_top_pt, x1 + img_width_pt, img_top_pt + img_height_pt))

            if _is_visual_only_label(label) and not rendered_visual:
                crop_bytes = _render_crop_bytes(src_page, fitz.Rect(x1, y1, x2, y2))
                if crop_bytes:
                    img_width_pt = max(x2 - x1, 2.0)
                    img_height_pt = max(y2 - y1, 2.0)
                    img_top_pt = _resolve_overlap_top(
                        left_pt=x1,
                        width_pt=img_width_pt,
                        top_pt=y1,
                        height_pt=img_height_pt,
                        placed=occupied_zones,
                        page_height=rect.height,
                        margin=0.4,
                        max_shift=None,
                    )
                    add_image_box(
                        paragraph,
                        image_bytes=crop_bytes,
                        left_pt=x1,
                        top_pt=img_top_pt,
                        width_pt=img_width_pt,
                        height_pt=img_height_pt,
                        z_index=image_z_index,
                    )
                    occupied_zones.append((x1, img_top_pt, x1 + img_width_pt, img_top_pt + img_height_pt))
                    image_z_index += 1
                    rendered_visual = True
                    if normalized_label in IMAGE_LIKE_LABELS:
                        rendered_image_like_bboxes.append((x1, img_top_pt, x1 + img_width_pt, img_top_pt + img_height_pt))

            text = _restore_form_layout((block.text or "").strip())
            if _is_visual_only_label(label):
                continue
            if not text:
                continue

            translated = translator.translate_text(text)
            translated, heading_level = _extract_markdown_heading(translated)
            underline_link = _looks_like_link_text(translated)
            is_bullet_list = _looks_like_bullet_list_text(translated)
            width_pt = _limit_text_box_width(
                block_index=block_index,
                mapped_blocks=mapped_blocks,
                column_regions=column_regions,
                page_width_pt=rect.width,
                desired_width_pt=width_pt,
            )
            left_pt = x1
            top_pt = y1
            left_pt, draw_width_pt, placed_bucket = _apply_column_bounds(
                current=mapped,
                desired_left_pt=left_pt,
                desired_width_pt=width_pt,
                column_regions=column_regions,
                page_width_pt=rect.width,
            )
            draw_width_pt, _ = _engine_textbox_safe_width(
                engine=engine,
                left_pt=left_pt,
                width_pt=draw_width_pt,
                page_width_pt=rect.width,
                bucket=placed_bucket,
            )
            if engine == "mineru" and _is_textbox_label(label):
                min_width = rect.width * 0.72
                max_width = max(2.0, rect.width - left_pt - 4.0)
                draw_width_pt = min(max_width, max(draw_width_pt, min_width))

            render_text = translated
            if engine in {"dotsocr", "paddleocr-vl", "mineru"}:
                if not _is_compact_heading_text(translated) and not is_bullet_list:
                    forced_lines = _wrap_text_to_lines(
                        text=translated,
                        width_pt=draw_width_pt,
                        font_size_pt=max(9.0, float(font_size)),
                        padding_pt=float(adaptive_profile["textBoxPadding"]),
                    )
                    if len(forced_lines) > 1:
                        render_text = "\n".join(forced_lines)

            effective_min_font_size = max(min_font_size, float(adaptive_profile["minFontSize"]))
            effective_default_font_size = max(6.5, float(font_size))
            if adaptive_profile["receiptMode"]:
                effective_default_font_size = max(effective_default_font_size, 7.2)
            else:
                effective_default_font_size = max(effective_default_font_size, 8.0)
            if engine == "dotsocr":
                effective_default_font_size = max(effective_default_font_size, 7.8)
            elif engine == "paddleocr-vl":
                effective_default_font_size = max(effective_default_font_size, 7.6)
            elif engine == "mineru":
                effective_default_font_size = max(effective_default_font_size, 8.0)
                effective_min_font_size = max(effective_min_font_size, 7.0)

            label_for_font = label
            effective_heading_level = heading_level if _is_title_like_label(label) else 0
            is_heading_style = effective_heading_level > 0 or _is_title_like_label(label)
            if is_heading_style:
                label_for_font = "paragraph_title"
                heading_target = min(10.0, float(font_size) + 1.0)
                effective_default_font_size = max(effective_default_font_size, heading_target)
                effective_default_font_size = min(effective_default_font_size, 10.0)
                effective_min_font_size = min(effective_min_font_size, max(min_font_size, float(font_size)))
            if engine in {"dotsocr", "paddleocr-vl"} and not tiny_box:
                font_size_pt = max(effective_min_font_size, effective_default_font_size)
            else:
                font_size_pt = _choose_font_size_pt(
                    text=render_text,
                    label=label_for_font,
                    width_pt=draw_width_pt,
                    height_pt=base_height_pt,
                    default_font_size=effective_default_font_size,
                    min_font_size=effective_min_font_size,
                    line_height=float(adaptive_profile["lineHeight"]),
                    padding_pt=float(adaptive_profile["textBoxPadding"]),
                )
            draw_height_pt = _expand_text_box_height(
                text=render_text,
                width_pt=draw_width_pt,
                base_height_pt=base_height_pt,
                font_size_pt=font_size_pt,
                line_height=float(adaptive_profile["lineHeight"]),
                padding_pt=float(adaptive_profile["textBoxPadding"]),
                page_height_pt=rect.height,
                top_pt=top_pt,
            )

            if tiny_box:
                left_pt, top_pt, draw_width_pt, draw_height_pt = _build_tiny_box(
                    x1=x1,
                    y1=y1,
                    width_pt=width_pt,
                    height_pt=max(font_size_pt * 1.15 + 1.0, 2.0),
                    page_rect=rect,
                    text=render_text,
                    font_size_pt=max(effective_min_font_size, font_size_pt),
                )
                font_size_pt = max(effective_min_font_size, min(font_size_pt, 3.0))
                draw_width_pt, _ = _engine_textbox_safe_width(
                    engine=engine,
                    left_pt=left_pt,
                    width_pt=draw_width_pt,
                    page_width_pt=rect.width,
                    bucket=placed_bucket,
                )
            else:
                if engine in {"dotsocr", "paddleocr-vl"}:
                    font_size_pt = max(effective_min_font_size, effective_default_font_size)
                else:
                    font_size_pt = _choose_font_size_pt(
                        text=render_text,
                        label=label_for_font,
                        width_pt=draw_width_pt,
                        height_pt=draw_height_pt,
                        default_font_size=effective_default_font_size,
                        min_font_size=effective_min_font_size,
                        line_height=float(adaptive_profile["lineHeight"]),
                        padding_pt=float(adaptive_profile["textBoxPadding"]),
                    )
                draw_height_pt = _expand_text_box_height(
                    text=render_text,
                    width_pt=draw_width_pt,
                    base_height_pt=draw_height_pt,
                    font_size_pt=font_size_pt,
                    line_height=float(adaptive_profile["lineHeight"]),
                    padding_pt=float(adaptive_profile["textBoxPadding"]),
                    page_height_pt=rect.height,
                    top_pt=top_pt,
                )

            placed_boxes = placed_boxes_map.setdefault(placed_bucket, [])
            small_box = draw_width_pt <= 120.0 or draw_height_pt <= 22.0
            if engine in {"dotsocr", "paddleocr-vl"}:
                overlap_margin = 0.25
            elif engine == "mineru":
                overlap_margin = 0.2
            else:
                overlap_margin = 1.2 if small_box else 0.3
            overlap_shift = None

            # 通用防越界兜底：若文本在当前框估算仍放不下，统一走分行行盒渲染。
            compact_heading = _is_compact_heading_text(render_text)
            need_force_split = (not compact_heading) and (not is_bullet_list) and (not _docx_text_fits(
                text=render_text,
                width_pt=draw_width_pt,
                height_pt=draw_height_pt,
                font_size_pt=font_size_pt,
                line_height=float(adaptive_profile["lineHeight"]),
                padding_pt=float(adaptive_profile["textBoxPadding"]),
            ))
            if ((not is_bullet_list) and _should_split_into_line_boxes(
                engine=engine,
                label=label,
                text=render_text,
                width_pt=draw_width_pt,
                bucket=placed_bucket,
            )) or need_force_split:
                split_lines = _wrap_text_to_lines(
                    text=render_text,
                    width_pt=draw_width_pt,
                    font_size_pt=font_size_pt,
                    padding_pt=float(adaptive_profile["textBoxPadding"]),
                )
                if len(split_lines) > 1:
                    line_extra_pt = 0.25 if engine in {"dotsocr", "paddleocr-vl", "mineru"} else _WORD_PARA_SPACING_EXTRA_PT
                    line_box_height = max(
                        font_size_pt * float(adaptive_profile["lineHeight"]) + line_extra_pt + 1.5,
                        6.0,
                    )
                    block_height = line_box_height * len(split_lines)
                    line_top = _resolve_overlap_top(
                        left_pt=left_pt,
                        width_pt=draw_width_pt,
                        top_pt=top_pt,
                        height_pt=min(block_height, rect.height),
                        placed=occupied_zones,
                        page_height=rect.height,
                        margin=overlap_margin,
                        max_shift=overlap_shift,
                    )
                    available_for_block = max(2.0, rect.height - line_top)
                    font_size_pt = _shrink_font_size_to_fit_height(
                        text=render_text,
                        width_pt=draw_width_pt,
                        available_height_pt=available_for_block,
                        current_font_size=font_size_pt,
                        min_font_size=effective_min_font_size,
                        line_height=float(adaptive_profile["lineHeight"]),
                        padding_pt=float(adaptive_profile["textBoxPadding"]),
                    )
                    split_lines = _wrap_text_to_lines(
                        text=render_text,
                        width_pt=draw_width_pt,
                        font_size_pt=font_size_pt,
                        padding_pt=float(adaptive_profile["textBoxPadding"]),
                    )
                    line_box_height = max(
                        font_size_pt * float(adaptive_profile["lineHeight"]) + line_extra_pt + 1.0,
                        5.0,
                    )
                    for line_text in split_lines:
                        current_text = line_text if line_text.strip() else " "
                        current_height = min(line_box_height, max(1.0, rect.height - line_top))
                        line_top = _resolve_overlap_top(
                            left_pt=left_pt,
                            width_pt=draw_width_pt,
                            top_pt=line_top,
                            height_pt=current_height,
                            placed=occupied_zones,
                            page_height=rect.height,
                            margin=overlap_margin,
                            max_shift=overlap_shift,
                        )
                        add_textbox(
                            paragraph,
                            text=current_text,
                            left_pt=left_pt,
                            top_pt=line_top,
                            width_pt=draw_width_pt,
                            height_pt=current_height,
                            font_name=font_name,
                            font_size_pt=font_size_pt,
                            z_index=text_z_index,
                            bold=is_heading_style,
                            underline=underline_link,
                        )
                        placed_box = (left_pt, line_top, left_pt + draw_width_pt, line_top + current_height)
                        placed_boxes.append(placed_box)
                        placed_boxes_all.append(placed_box)
                        occupied_zones.append(placed_box)
                        text_z_index += 1
                        rendered_blocks += 1
                        line_top += current_height
                    continue

            # 碰撞检测：如果当前框与已放置框有垂直重叠，向下推移
            top_pt = _resolve_overlap_top(
                left_pt=left_pt,
                width_pt=draw_width_pt,
                top_pt=top_pt,
                height_pt=draw_height_pt,
                placed=occupied_zones,
                page_height=rect.height,
                margin=overlap_margin,
                max_shift=overlap_shift,
            )

            available_height = max(2.0, rect.height - top_pt)
            font_size_pt = _shrink_font_size_to_fit_height(
                text=render_text,
                width_pt=draw_width_pt,
                available_height_pt=available_height,
                current_font_size=font_size_pt,
                min_font_size=effective_min_font_size,
                line_height=float(adaptive_profile["lineHeight"]),
                padding_pt=float(adaptive_profile["textBoxPadding"]),
            )
            draw_height_pt = _expand_text_box_height(
                text=render_text,
                width_pt=draw_width_pt,
                base_height_pt=draw_height_pt,
                font_size_pt=font_size_pt,
                line_height=float(adaptive_profile["lineHeight"]),
                padding_pt=float(adaptive_profile["textBoxPadding"]),
                page_height_pt=rect.height,
                top_pt=top_pt,
            )

            add_textbox(
                paragraph,
                text=render_text,
                left_pt=left_pt,
                top_pt=top_pt,
                width_pt=draw_width_pt,
                height_pt=draw_height_pt,
                font_name=font_name,
                font_size_pt=font_size_pt,
                z_index=text_z_index,
                bold=is_heading_style,
                underline=underline_link,
            )
            placed_box = (left_pt, top_pt, left_pt + draw_width_pt, top_pt + draw_height_pt)
            placed_boxes.append(placed_box)
            placed_boxes_all.append(placed_box)
            occupied_zones.append(placed_box)
            text_z_index += 1
            rendered_blocks += 1

        rendered_pages += 1
        flow_cursor_pt = 0.0
        for table_item in sorted(pending_tables, key=lambda item: float(item.get("top", 0.0))):
            top_pt = max(0.0, float(table_item.get("top", 0.0)))
            if top_pt > flow_cursor_pt:
                _add_vertical_spacer(doc, top_pt - flow_cursor_pt)
            _render_word_table(doc, table_item["rows"], font_name=font_name, font_size_pt=font_size)
            doc.add_paragraph()
            flow_cursor_pt = max(flow_cursor_pt, top_pt + max(2.0, float(table_item.get("height", 2.0))))

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    src_doc.close()
    return {
        "outputDocxPath": str(output_docx),
        "providerKind": provider_kind,
        "pages": rendered_pages,
        "textBlocks": rendered_blocks,
    }


def main() -> None:
    args = _parse_args()
    input_dir = Path(args.input_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    search_dirs = [Path(p).expanduser().resolve() for p in args.pdf_search_dir]
    translator = Translator(
        enabled=args.enable_translate,
        source_lang=args.source_lang,
        target_lang=args.target_lang,
    )

    matches = _iter_raw_json_files(input_dir)
    if not matches:
        raise FileNotFoundError(f"No *_<engine>_raw.json files found in: {input_dir}")

    results: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    for raw_json_path, sample_name, engine in matches:
        try:
            payload = json.loads(raw_json_path.read_text(encoding="utf-8"))
            input_pdf = _resolve_input_pdf(sample_name, search_dirs)
            output_docx = output_dir / f"{sample_name}_{engine}_visual_clone.docx"
            summary = _export_single_docx(
                payload=payload,
                engine=engine,
                raw_json_path=raw_json_path,
                input_pdf=input_pdf,
                output_docx=output_docx,
                geometry_source=args.geometry_source,
                with_page_background=args.with_page_background,
                font_name=args.font_name,
                font_size=args.font_size,
                min_font_size=args.min_font_size,
                text_box_extra_width_pt=args.text_box_extra_width_pt,
                text_box_extra_height_pt=args.text_box_extra_height_pt,
                translator=translator,
            )
            results.append(
                {
                    "sample": sample_name,
                    "engine": engine,
                    "rawJsonPath": str(raw_json_path),
                    "inputPdfPath": str(input_pdf),
                    **summary,
                }
            )
            print(f"[OK] {sample_name} [{engine}] -> {output_docx}")
        except Exception as exc:  # noqa: BLE001
            errors.append({"sample": sample_name, "engine": engine, "rawJsonPath": str(raw_json_path), "error": str(exc)})
            print(f"[ERROR] {sample_name} [{engine}] -> {exc}")

    output_dir.mkdir(parents=True, exist_ok=True)
    manifest = {
        "inputDir": str(input_dir),
        "outputDir": str(output_dir),
        "geometrySource": args.geometry_source,
        "withPageBackground": args.with_page_background,
        "results": results,
        "errors": errors,
    }
    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Saved manifest: {manifest_path}")

    if errors and not results:
        raise RuntimeError(f"All exports failed: {errors}")


if __name__ == "__main__":
    main()
